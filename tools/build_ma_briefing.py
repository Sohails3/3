"""
Builds MA_Briefing.docx — a condensed M&A / IB briefing document.
Run: .venv/bin/python tools/build_ma_briefing.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── GP Bullhound colours ──────────────────────────────────────────────────────
NAVY   = RGBColor(0x25, 0x28, 0x50)   # Night Blue
RED    = RGBColor(0xCC, 0x06, 0x05)   # Traffic Red
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREY   = RGBColor(0xF4, 0xF4, 0xF4)
DARK   = RGBColor(0x2C, 0x2C, 0x2C)
MUTED  = RGBColor(0x6B, 0x72, 0x80)

OUT = Path(__file__).parent.parent / "MA_Briefing.docx"


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, side="left", color="CC0605", sz=18):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"),   "single")
    border.set(qn("w:sz"),    str(sz))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)
    tcBorders.append(border)

def no_space(para):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)

def add_para(doc, text="", bold=False, size=10, color=DARK, align=None,
             space_before=0, space_after=4, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = color
    return p

def add_inline(para, text, bold=False, italic=False, size=10, color=DARK):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    return run

def section_header(doc, number, title):
    """Red number tag + navy title in one paragraph."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, NAVY)
    set_cell_border(cell, "left",   "CC0605", 24)
    set_cell_border(cell, "top",    "252850", 4)
    set_cell_border(cell, "right",  "252850", 4)
    set_cell_border(cell, "bottom", "252850", 4)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.2)
    r1 = p.add_run(f"{number}  ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = RED
    r2 = p.add_run(title.upper())
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = WHITE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def sub_header(doc, text):
    p = add_para(doc, text.upper(), bold=True, size=8, color=RED,
                 space_before=8, space_after=2)
    # underline via border on paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CC0605")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def bullet(doc, text, bold_prefix=None, indent=0.4):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Cm(indent)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    if bold_prefix:
        add_inline(p, bold_prefix + " ", bold=True, size=9.5, color=DARK)
        add_inline(p, text, size=9.5, color=DARK)
    else:
        add_inline(p, text, size=9.5, color=DARK)

def add_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_bg(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(h.upper())
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = WHITE

    # data rows
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = GREY if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            if isinstance(val, tuple):
                text, bold = val
            else:
                text, bold = val, False
            r = p.add_run(str(text))
            r.bold = bold
            r.font.size = Pt(9)
            r.font.color.rgb = DARK

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


# ── Build document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# Default paragraph style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── COVER ─────────────────────────────────────────────────────────────────────

# Red top bar (1-row, 1-col table)
bar = doc.add_table(rows=1, cols=1)
bar.alignment = WD_TABLE_ALIGNMENT.LEFT
bc = bar.cell(0, 0)
set_cell_bg(bc, RED)
no_space(bc.paragraphs[0])
bc.paragraphs[0].paragraph_format.space_before = Pt(3)
bc.paragraphs[0].paragraph_format.space_after  = Pt(3)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_para(doc, "GP BULLHOUND", bold=True, size=11, color=RED,
         space_before=0, space_after=0)
add_para(doc, "M&A Advisory", bold=False, size=9, color=MUTED,
         space_before=0, space_after=14)

p_title = add_para(doc, "M&A & Investment Banking", bold=True, size=22,
                   color=NAVY, space_before=0, space_after=4)
add_para(doc, "Briefing Document", bold=False, size=16, color=MUTED,
         space_before=0, space_after=2)
add_para(doc, "Strategic Fit Engine  ·  Confidential", italic=True,
         size=9, color=MUTED, space_before=0, space_after=0)

# Divider
div = doc.add_table(rows=1, cols=1)
div.alignment = WD_TABLE_ALIGNMENT.LEFT
dc = div.cell(0, 0)
set_cell_bg(dc, RED)
no_space(dc.paragraphs[0])
dc.paragraphs[0].paragraph_format.space_before = Pt(2)
dc.paragraphs[0].paragraph_format.space_after  = Pt(2)
doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── SECTION 1: COMPANY METRICS ────────────────────────────────────────────────

section_header(doc, "01", "Company Metrics")

add_para(doc, "The five numbers that anchor every deal conversation.", italic=True,
         size=9, color=MUTED, space_before=0, space_after=6)

metrics = [
    ("ARR",  "Annual Recurring Revenue",
     "Annualised value of active subscription contracts. The valuation anchor for SaaS.",
     "Rule: ARR × Multiple = Enterprise Value. Growing ARR at 40%+ commands premium multiples."),
    ("EBITDA", "Earnings Before Interest, Tax, Depreciation & Amortisation",
     "Operating profitability, stripped of financing and accounting noise. Proxy for cash generation.",
     "Healthy SaaS margin at scale: 20–30%. Early stage is often negative — that's expected."),
    ("FCF",  "Free Cash Flow",
     "Cash left after capex. What the DCF model actually discounts.",
     "FCF = EBITDA × Conversion Rate. SaaS is capital-light so conversion is typically 50–70%."),
    ("NRR",  "Net Revenue Retention",
     "% of last year's ARR retained and expanded, excluding new customers.",
     ">100% = existing customers are growing spend. >120% = best-in-class. <90% = red flag."),
    ("Burn Rate", "Monthly cash outflow above income",
     "Relevant for pre-profit targets. Runway = Cash ÷ Monthly Burn.",
     "<12 months runway = urgency to sell. Creates negotiating leverage for the buyer."),
]

for short, full, what, why in metrics:
    sub_header(doc, f"{short} — {full}")
    bullet(doc, what)
    bullet(doc, why, bold_prefix="Why it matters:")

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── SECTION 2: VALUATION ─────────────────────────────────────────────────────

section_header(doc, "02", "Valuation")

sub_header(doc, "Revenue Multiple  (EV / ARR)")
add_para(doc,
    "The primary valuation shorthand for high-growth software. "
    "Enterprise Value = ARR × Multiple. The multiple is driven by growth rate, NRR, "
    "gross margin, and market size.",
    size=9.5, space_before=2, space_after=4)

add_table(doc,
    ["ARR Growth Rate", "Typical EV / ARR Range"],
    [
        ("> 60% YoY",  "8 – 15×"),
        ("40 – 60% YoY", "5 – 8×"),
        ("20 – 40% YoY", "3 – 6×"),
        ("< 20% YoY",  "1.5 – 3×"),
    ],
    col_widths=[7.5, 7.5]
)

sub_header(doc, "EBITDA Multiple  (EV / EBITDA)")
add_para(doc,
    "Dominant metric for profitable businesses. More meaningful than revenue multiples "
    "at maturity. European healthcare IT: typically 15–25×.",
    size=9.5, space_before=2, space_after=6)

sub_header(doc, "DCF — Discounted Cash Flow")
add_para(doc,
    "The most rigorous method. Values all future cash flows in today's money. "
    "Every figure in the DCF section of the report flows from these five inputs:",
    size=9.5, space_before=2, space_after=4)

add_table(doc,
    ["Input", "Default", "What It Does"],
    [
        ("Revenue Growth",     "10% → 6% declining",  "Projects top-line over 5 years. Declining rate reflects market maturation."),
        ("EBITDA Margin",      "15%",                  "Converts revenue to operating profit."),
        ("FCF Conversion",     "55%",                  "Converts EBITDA to cash after capex and working capital."),
        ("WACC",               "10.5%",                "The discount rate. Higher = lower valuation. Reflects risk and cost of capital."),
        ("Terminal Growth",    "3.0%",                 "Long-run growth rate after year 5. Anchored to nominal GDP. Drives 60–80% of EV."),
    ],
    col_widths=[3.5, 3.5, 9.5]
)

sub_header(doc, "Reading the DCF Output")
add_table(doc,
    ["Line", "Meaning"],
    [
        ("Sum of PV FCFs",      "Discounted value of years 1–5 cash flows"),
        ("PV of Terminal Value","All cash flows after year 5, discounted to today"),
        ("Enterprise Value",    "Sum of above — standalone intrinsic value"),
        ("Synergy Value",       "Extra value this specific buyer can unlock"),
        ("Acquisition Price",   "EV + Synergy — indicative offer price"),
        ("EV / Revenue",        "Implied multiple — cross-check against comparables"),
    ],
    col_widths=[5.0, 10.5]
)

sub_header(doc, "Synergy Premium")
add_para(doc,
    "The extra value a specific acquiror can generate that a standalone buyer cannot. "
    "Our model applies 20% on EV. Three types:",
    size=9.5, space_before=2, space_after=4)
bullet(doc, "Cross-sell the target's product to the buyer's existing customers",
       bold_prefix="Revenue:")
bullet(doc, "Shared infrastructure, combined G&A, reduced headcount overlap",
       bold_prefix="Cost:")
bullet(doc, "Target's IP or data enhances the buyer's core product",
       bold_prefix="Capability:")
add_para(doc, "Note: 50–70% of deals fail to deliver projected synergies. "
         "Always stress-test to zero synergies to confirm the deal works on a standalone basis.",
         italic=True, size=8.5, color=MUTED, space_before=4, space_after=6)


# ── SECTION 3: DEAL PROCESS ───────────────────────────────────────────────────

section_header(doc, "03", "Deal Process")

sub_header(doc, "Stages")
stages = [
    ("1. Origination",    "Identify target, build thesis"),
    ("2. Approach",       "NDA signed, initial discussions"),
    ("3. Indicative LOI", "Non-binding letter of intent — price & key terms"),
    ("4. Due Diligence",  "Financial, legal, technical, commercial DD — typically 6–10 weeks"),
    ("5. Binding Offer",  "Final price and conditions"),
    ("6. SPA Signed",     "Share Purchase Agreement — legally binding"),
    ("7. Regulatory",     "Competition authority clearance if required"),
    ("8. Close",          "Cash transferred, ownership transferred"),
    ("9. Integration",    "100-day plan: people, systems, culture"),
]
for stage, desc in stages:
    bullet(doc, desc, bold_prefix=stage)

sub_header(doc, "Key Documents")
add_table(doc,
    ["Document", "What It Is"],
    [
        ("NDA",  "Non-Disclosure Agreement. Signed before any confidential data is shared."),
        ("CIM",  "Confidential Information Memorandum. 30–80 page sell-side document: financials, product, market, team, valuation."),
        ("LOI",  "Letter of Intent. Non-binding offer. Triggers exclusivity (usually 30–60 days)."),
        ("Data Room", "Secure repository of all DD materials: accounts, contracts, IP, cap table."),
        ("SPA",  "Share Purchase Agreement. The binding contract. Contains reps, warranties, indemnities."),
    ],
    col_widths=[3.0, 12.5]
)

sub_header(doc, "Deal Structures — How Consideration Is Paid")
add_table(doc,
    ["Type", "Description", "Seller Implication"],
    [
        ("Cash",     "Full payment at close",               "Certainty. Taxable immediately."),
        ("Stock",    "Buyer shares issued to seller",       "Upside if buyer performs. Dilution risk."),
        ("Earnout",  "Deferred cash tied to milestones",    "Aligns incentives. Often disputed post-close."),
        ("Rollover", "Seller reinvests % into combined co.", "Retained upside. Common in PE transactions."),
    ],
    col_widths=[3.0, 6.5, 6.5]
)


# ── SECTION 4: STRATEGIC FIT SCORING ─────────────────────────────────────────

section_header(doc, "04", "Strategic Fit Scoring")

add_para(doc,
    "The 7-criterion rubric forces explicit, pre-committed criteria before data is seen — "
    "reducing the anchoring bias that distorts most deal processes.",
    size=9.5, italic=True, space_before=2, space_after=6)

sub_header(doc, "The 7 Criteria")
add_table(doc,
    ["ID", "Criterion", "Type"],
    [
        ("C1–C4", "Buyer-specific (derived from acquiror's acquisition history and stated strategy)", "Custom per buyer"),
        ("C5",    "Regional & Operational Fit — geography, cloud architecture, integration cost",     "Universal"),
        ("C6",    "Revenue Synergy Potential — cross-sell to acquiror's customer base",               "Universal"),
        ("C7",    "Ease of Acquisition — ownership, regulatory risk, competitive tension",            "Universal"),
    ],
    col_widths=[2.0, 11.0, 3.5]
)

sub_header(doc, "Score Interpretation")
add_table(doc,
    ["Score", "Meaning", "Tier"],
    [
        ("5",     "Direct, traceable match to buyer's documented gap",      "Tier 1 — Proceed (30–35/35)"),
        ("3–4",   "Partial alignment; meaningful but not perfect fit",      "Tier 2 — Monitor (22–29/35)"),
        ("1–2",   "Friction point; multiple low scores = flag for review",  "Tier 3 — Discard (<22/35)"),
    ],
    col_widths=[1.5, 9.5, 5.5]
)

sub_header(doc, "Common Deal-Breaker Risks")
risks = [
    ("IP ownership dispute", "You may not be acquiring what you think you're buying"),
    ("Key-man dependency",   "Business value leaves with the founder"),
    ("Revenue concentration","One customer >30% of ARR that may churn post-deal"),
    ("Regulatory approval",  "Sector regulators may block change of control"),
    ("Competitive bid",      "Another well-funded buyer forces an irrational price"),
]
for risk, why in risks:
    bullet(doc, why, bold_prefix=f"{risk}:")


# ── SECTION 5: HOW TO TALK THROUGH THE REPORT ────────────────────────────────

section_header(doc, "05", "Talking Points — Meeting Script")

sub_header(doc, "Open (0–2 min)")
add_para(doc,
    '"This analysis screens [N] private companies in [sector] across [geography] against a '
    '7-criterion rubric derived from [Buyer]\'s own acquisition history. We ran a full DCF '
    'on the top-ranked target. I\'ll lead with the recommendation."',
    italic=True, size=9.5, space_before=2, space_after=6)

sub_header(doc, "Lead with the Recommendation (2–5 min)")
add_para(doc,
    '"Our primary recommendation is [Company]. It scored [X/35] — the highest in the cohort — '
    'driven by [top 2 criteria]. It is the only company that directly addresses [buyer\'s gap]. '
    'Indicative acquisition price including a 20% synergy premium is €[X]M, implying [Y]× ARR."',
    italic=True, size=9.5, space_before=2, space_after=6)

sub_header(doc, "Defending the DCF (if challenged)")
qa = [
    ("Why is terminal value so large?",
     '"Structurally inherent to DCF on high-growth businesses. The terminal value uses 3% '
     'perpetuity growth — below GDP — deliberately conservative. It is not speculative."'),
    ("How reliable is the financial data?",
     '"Indicative, based on public filings and sector benchmarks. The data room resolves '
     'material uncertainties. That is exactly why we recommend an NDA approach first."'),
    ("Is the 20% synergy premium justifiable?",
     '"Stress-tested to zero synergies — the deal is still accretive at [X]× ARR on a '
     'standalone basis. The premium is what makes it worth paying above fair value."'),
    ("What if founders won\'t sell?",
     '"The C7 score reflects our view on seller motivation. Next step is a discreet approach '
     'via [mutual investor / banker]. Hostile approaches in private M&A destroy goodwill."'),
]
for q, a in qa:
    p = add_para(doc, space_before=3, space_after=2)
    add_inline(p, f"Q: {q}", bold=True, size=9, color=NAVY)
    add_para(doc, a, italic=True, size=9, color=DARK, space_before=0, space_after=5)


# ── SECTION 6: GLOSSARY ───────────────────────────────────────────────────────

section_header(doc, "06", "Key Terms Glossary")

terms = [
    ("ARR",         "Annual Recurring Revenue — annualised subscription contract value"),
    ("NRR",         "Net Revenue Retention — % of ARR retained + expanded from existing customers"),
    ("CAC / LTV",   "Customer Acquisition Cost / Lifetime Value. LTV/CAC >3× = healthy unit economics"),
    ("Rule of 40",  "Growth % + EBITDA Margin % ≥ 40 — benchmark for SaaS health"),
    ("WACC",        "Weighted Average Cost of Capital — the DCF discount rate"),
    ("Terminal Value","Present value of all cash flows beyond the 5-year forecast"),
    ("EV",          "Enterprise Value — equity + debt − cash. Total business value"),
    ("Equity Value","EV minus net debt. What shareholders actually receive"),
    ("EBITDA",      "Earnings Before Interest, Tax, Depreciation & Amortisation"),
    ("FCF",         "Free Cash Flow — cash after capex, available to investors"),
    ("Cap Table",   "Capitalisation table — who owns what % of the company"),
    ("LOI",         "Letter of Intent — non-binding statement of deal terms"),
    ("SPA",         "Share Purchase Agreement — the binding sale contract"),
    ("QoE",         "Quality of Earnings — independent verification of financials"),
    ("Earnout",     "Deferred consideration tied to post-close performance milestones"),
    ("Dry Powder",  "Undeployed capital held by a fund, ready to invest"),
    ("Buy-side",    "Advisory mandate for the acquiror (the buyer)"),
    ("Sell-side",   "Advisory mandate for the company being sold"),
    ("Strategic Buyer","A company buying for operational/product synergies"),
    ("Financial Buyer","A private equity fund buying for financial return"),
    ("Control Premium","Extra price paid to acquire a majority controlling stake"),
]

rows = [(t, d) for t, d in terms]
add_table(doc, ["Term", "Definition"], rows, col_widths=[4.0, 12.5])


# ── FOOTER ────────────────────────────────────────────────────────────────────

doc.add_paragraph().paragraph_format.space_after = Pt(8)
footer_bar = doc.add_table(rows=1, cols=1)
footer_bar.alignment = WD_TABLE_ALIGNMENT.LEFT
fc = footer_bar.cell(0, 0)
set_cell_bg(fc, NAVY)
p = fc.paragraphs[0]
p.paragraph_format.space_before = Pt(5)
p.paragraph_format.space_after  = Pt(5)
p.paragraph_format.left_indent  = Cm(0.3)
r = p.add_run("GP Bullhound  ·  Technology Investment Banking  ·  Strictly Confidential  ·  Not for Distribution")
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xCC)

doc.save(OUT)
print(f"Saved: {OUT}")
